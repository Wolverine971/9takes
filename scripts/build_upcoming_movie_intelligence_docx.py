#!/usr/bin/env python3
# scripts/build_upcoming_movie_intelligence_docx.py
"""Build the cited 9takes upcoming-movie intelligence DOCX."""

from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/content-research/2026-08-29_upcoming-movie-intelligence-for-9takes.md"
OUTPUT = ROOT / "docs/content-research/2026-08-29_upcoming-movie-intelligence-for-9takes.docx"

NAVY = RGBColor(32, 55, 72)
INK_BLUE = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(181, 133, 43)
GRAY = RGBColor(80, 80, 80)
MUTED = RGBColor(105, 112, 120)
WHITE = RGBColor(255, 255, 255)
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF8E8"
PALE_BLUE = "EDF3F8"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_dxa / 1440)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D7DBE2", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)
        tag.set(qn("w:space"), "0")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    r_pr.extend([r_fonts, color, underline, sz])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|\*[^*]+\*|https?://\S+)"
)


def add_inline(paragraph, text, *, base_size=11, base_color=None, base_bold=False):
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=base_size, color=base_color, bold=base_bold)
        token = match.group(0)
        link_match = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
        if link_match:
            add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, color=base_color, bold=base_bold, italic=True)
        else:
            clean_url = token.rstrip(".,;)")
            add_hyperlink(paragraph, "Open source", clean_url)
            trailing = token[len(clean_url) :]
            if trailing:
                run = paragraph.add_run(trailing)
                set_run_font(run, size=base_size, color=base_color, bold=base_bold)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color, bold=base_bold)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.keep_with_next = False
    normal.paragraph_format.keep_together = False
    normal.paragraph_format.page_break_before = False

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.keep_with_next = False
        style.paragraph_format.keep_together = False
        style.paragraph_format.page_break_before = False


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_page_number_start(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def configure_section(section, *, content=False):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    if not content:
        return

    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        run = hp.add_run("9takes  |  Upcoming Movie Intelligence")
        set_run_font(run, size=9, color=MUTED, bold=True)

    for footer in (section.footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        run = fp.add_run("Page ")
        set_run_font(run, size=9, color=MUTED)
        add_field(fp, "PAGE")
    set_page_number_start(section, 1)


def add_cover(doc):
    section = doc.sections[0]
    configure_section(section, content=False)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("EDITORIAL INTELLIGENCE REPORT")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Upcoming Movie Intelligence")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    run = subtitle.add_run("What 9takes should publish before")
    set_run_font(run, size=15, color=DARK_BLUE)
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.paragraph_format.space_after = Pt(28)
    run = subtitle2.add_run("the fall 2026 and 2027 attention waves")
    set_run_font(run, size=15, color=DARK_BLUE)

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(72)
    run = descriptor.add_run("Breakout films  |  emerging actors  |  psychology themes  |  publishing calendar")
    set_run_font(run, size=10.5, color=GOLD, italic=True)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(4)
    run = date.add_run("August 29, 2026")
    set_run_font(run, size=12, color=NAVY, bold=True)
    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audience.paragraph_format.space_after = Pt(0)
    run = audience.add_run("Prepared for DJ and the 9takes editorial pipeline")
    set_run_font(run, size=9.5, color=GRAY, italic=True)


def add_lead_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color="D9C38B", size="8")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_GOLD)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_inline(
        p,
        "DECISION: Publish the existing Jeremy Allen White asset, then prioritize Austin Abrams, Tom Rhys Harries, and Joseph Zada. Build the cross-film editorial lane around epistemic anxiety and identity horror.",
        base_size=10.5,
        base_color=INK_BLUE,
        base_bold=True,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_at_a_glance_table(doc):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("Top opportunities at a glance")
    rows = [
        ("The Social Reckoning", "Very high", "Very high", "Existing assets", "Publish White; refresh Madison"),
        ("Clayface", "High", "Very high", "Tom Rhys Harries", "Profile by mid-September"),
        ("Sunrise on the Reaping", "Very high", "Very high", "Joseph Zada", "Profile by late September"),
        ("Resident Evil + Whalefall", "High", "High", "Austin Abrams", "Publish before September 18"),
        ("Klara and the Sun", "Med-high", "Very high", "Mia Tharia", "AI attachment piece; monitor Mia"),
        ("Verity", "High", "Very high", "None", "Unreliable-narrator article"),
        ("Hope", "Med-high", "High", "None", "Global-crossover preview"),
        ("Other Mommy", "Med-high", "High", "Child - no typing", "Attachment/doppelgänger essay"),
        ("Appofeniacs", "Low reach", "Very high", "None", "Evergreen deepfake/apophenia post"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [2448, 1080, 1080, 1800, 2952]
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["Film", "Reach", "Psych fit", "Actor opportunity", "Best action"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 3, 4) else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=9, color=INK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])

    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        cells = row.cells
        for cidx, text in enumerate(row_data):
            p = cells[cidx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx in (0, 3, 4) else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(text)
            set_run_font(run, size=8.6, color=INK_BLUE if cidx == 0 else None, bold=(cidx == 0))
            cells[cidx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ridx % 2 == 1:
                shade_cell(cells[cidx], "FAFBFC")
    set_table_geometry(table, widths)
    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    run = source.add_run("Editorial ranking; reach is a qualitative projection, not a box-office estimate.")
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_heading(doc, level, text):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_inline(p, text, base_size={1: 16, 2: 13, 3: 12}[level], base_color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], base_bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = False
    p.paragraph_format.page_break_before = False
    if text.startswith("**Projection:**"):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
    add_inline(p, text)
    return p


def normalize_text(text):
    return (
        text.replace("—", "-")
        .replace("–", "-")
        .replace("’", "'")
        .replace("“", '"')
        .replace("”", '"')
        .replace("…", "...")
        .replace("‑", "-")
    )


def parse_source(doc, bullet_num_id, decimal_num_id):
    lines = normalize_text(SOURCE.read_text(encoding="utf-8")).splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## Scope and assumptions"))
    inserted_callout = False
    inserted_matrix = False
    page_break_heads = {
        "Priority ranking: what to act on",
        "Secondary watchlist: meaningful but not first",
        "Actor priority list",
        "Trends 9takes can own",
        "Publishing calendar",
        "Claim-to-source ledger",
    }
    for raw in lines[start:]:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            text = line[3:].strip()
            heading = add_heading(doc, 1, text)
            if text in page_break_heads and len(doc.paragraphs) > 4:
                heading.paragraph_format.page_break_before = True
            if text == "Executive answer" and not inserted_callout:
                add_lead_callout(doc)
                inserted_callout = True
            if text == "Priority ranking: what to act on" and not inserted_matrix:
                add_at_a_glance_table(doc)
                inserted_matrix = True
        elif line.startswith("### "):
            text = line[4:].strip()
            heading = add_heading(doc, 2, text)
            if text.startswith(("5. Klara and the Sun", "12. Dune: Part Three")):
                heading.paragraph_format.page_break_before = True
        elif line.startswith("#### "):
            add_heading(doc, 3, line[5:].strip())
        elif re.match(r"^\d+\. \*\*", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, decimal_num_id)
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.keep_together = False
            add_inline(p, text)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            apply_numbering(p, bullet_num_id)
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.keep_together = False
            add_inline(p, line[2:])
        else:
            add_body(doc, line)


def audit_docx(path):
    with zipfile.ZipFile(path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
        numbering_xml = zf.read("word/numbering.xml").decode("utf-8")
    checks = {
        "letter_page": 'w:w="12240"' in document_xml and 'w:h="15840"' in document_xml,
        "one_inch_margins": 'w:top="1440"' in document_xml and 'w:left="1440"' in document_xml,
        "body_calibri": "Calibri" in styles_xml,
        "h1_blue": "2E74B5" in styles_xml,
        "real_numbering": "w:abstractNum" in numbering_xml and "w:numPr" in document_xml,
        "fixed_table_width": (
            'w:tblW w:w="9360" w:type="dxa"' in document_xml
            or 'w:tblW w:type="dxa" w:w="9360"' in document_xml
        ),
        "table_indent": (
            'w:tblInd w:w="120" w:type="dxa"' in document_xml
            or 'w:tblInd w:type="dxa" w:w="120"' in document_xml
        ),
        "no_fake_bullet_prefixes": "<w:t>- " not in document_xml,
    }
    failed = [key for key, value in checks.items() if not value]
    print("Preset audit:", checks)
    if failed:
        raise RuntimeError(f"DOCX preset audit failed: {failed}")


def build():
    doc = Document()
    doc.core_properties.title = "Upcoming Movie Intelligence for 9takes"
    doc.core_properties.subject = "Fall 2026 and 2027 editorial opportunities"
    doc.core_properties.author = "9takes editorial research"
    doc.core_properties.keywords = "movies, psychology, actors, trends, editorial calendar"
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    add_cover(doc)

    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(content_section, content=True)
    bullet_num_id = add_numbering(doc, "bullet")
    decimal_num_id = add_numbering(doc, "decimal")
    parse_source(doc, bullet_num_id, decimal_num_id)

    doc.settings.element.append(OxmlElement("w:updateFields"))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    audit_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise
